# -*- coding: utf-8 -*-
import datetime
import json
import jieba
import re
import jieba
import time
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn
key = ['暴雨', '大雪', '沙尘暴']
num1=1
num2=1
num3=1
keyrain = ['暴雨','中到大雨']
keysnow = ['大雪','中雪']
def change(line,str):
    part1 = re.split(r"[；，。（]", line)
    for tmp in part1:
        flag = 0
        word_list = jieba.lcut(tmp)
        for keyword in word_list:
            if keyword in key:
                flag = 1;
        if flag == 1:
            str += '，'
            str += tmp
    return str

def change2(day,snow,rain,wind):
    global num1,num2,num3
    part1 = re.split(r"[，。]", day)
    #print(part1)
    for tmp in part1:
        flag = 0
        word_list = jieba.lcut(tmp)
        #print(word_list)
        for keyword in word_list:
            if keyword in keyrain:
                flag=1
            elif keyword in keysnow:
                flag=2
            elif keyword=='风':
                flag=3
        if flag == 1:
            rain +='（'+str(num1)+'）'+tmp+'。'
            num1+=1
        if flag == 2:
            snow +='（'+str(num2)+'）'+tmp+'。'
            num2+=1
        if flag == 3:
            wind +='（'+str(num3)+'）'+tmp+'。'
            num3+=1
    return rain,snow,wind


publish ="    "+time.strftime('%d', time.localtime(time.time())) + "日"
with open('../weatherBriefingSpider/brief.json', 'r', encoding='utf-8') as fObj:
    raw_list = json.load(fObj)
    raw = raw_list[0]
    #完成发布部分内容
    line1 = raw['brief_detail']
    publish = change(line1, publish)
    publish += "。未来两者三日"
    line1 = raw['day1_detail']
    publish = change(line1, publish)
    publish += "，可能对交通运输产生影响。"
    #print(publish)

    #重要天气预报
    import_weather =""
    import_weather+=''.join(raw['key_point_title'])+'\n'
    import_weather+="    "+' '.join(raw['key_point_detail'])
    #print(import_weather)

    #图片
    img_urls=raw['img_urls']

    #未来三天具体天气预报
    future_weather="三、未来三天具体预报\n"
    day1=''.join(raw['day1_detail'])
    day2=''.join(raw['day2_detail'])
    day3=''.join(raw['day3_detail']) 
    #print(future_weather,day1,day2,day3)

    #雪、雨、风
    snow = ''
    rain = ''
    wind = ''
    #print(day1)
    rain,snow,wind=change2(day1,snow,rain,wind)
    rain,snow,wind=change2(day2,snow,rain,wind)
    rain,snow,wind=change2(day3,snow,rain,wind)
    #print(rain)
    #print(snow)
    #print(wind)

path='brief.docx'
if os.path.exists(path):  # 如果文件存在
    os.remove(path)  
fd = open(path, mode="w", encoding="utf-8")
fd.close()
document=Document()
# 发布
p1 = document.add_paragraph()
run=p1.add_run('    发布') 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.bold = True
run.font.color.rgb = RGBColor(0,0,0)
# 发布内容
p2 = document.add_paragraph()
run=p2.add_run(publish) 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.bold = True
run.font.color.rgb = RGBColor(0,0,0)
document.add_paragraph('')
document.add_paragraph('')
# 参考资料
p3 = document.add_paragraph()
run=p3.add_run("    "+"参考材料：http://www.nmc.cn/publish/weather-bulletin/index.htm") 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.bold = True
run.font.color.rgb = RGBColor(0,0,0)
document.add_paragraph('')
document.add_paragraph('')
# 重点天气预报
p4 = document.add_paragraph()
run=p4.add_run("    "+"二、重要天气预报") 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.bold = True
run.font.color.rgb = RGBColor(0,0,0)

p5 = document.add_paragraph()
run=p5.add_run("    "+import_weather) 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(0,0,0)
# 图片
s1 = document.add_paragraph()
run = s1.add_run("")
now = datetime.datetime.now()
date = str(now.year) + '年' + str(now.month) + '月' + str(now.day) + '日'
run.add_picture('../weatherBriefingSpider/weatherBriefingSpider/img/'+date+'-今日全国降水量预报图.jpg', width=Inches(1.85))
run.add_picture('../weatherBriefingSpider/weatherBriefingSpider/img/'+date+'-明日全国降水量预报图.jpg', width=Inches(1.85))
run.add_picture('../weatherBriefingSpider/weatherBriefingSpider/img/'+date+'-后天全国降水量预报图.jpg', width=Inches(1.85))
# 三、未来三日具体预报
p7 = document.add_paragraph()
run=p7.add_run("    "+"三、未来三日具体预报")
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.bold = True
run.font.color.rgb = RGBColor(0,0,0)

p8 = document.add_paragraph()
run=p8.add_run("    "+day1) 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(0,0,0)

p9 = document.add_paragraph()
run=p9.add_run("    "+day2) 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(0,0,0)

p10 = document.add_paragraph()
run=p10.add_run("    "+day3) 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(0,0,0)
# 雪
document.add_paragraph('')
p11 = document.add_paragraph()
run=p11.add_run("    "+"雪：") 
run.font.name = u'宋体'
run.font.size = Pt(12)
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(255,0,0)
run=p11.add_run(snow) 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(0,0,0)
#雨
p12 = document.add_paragraph()
run=p12.add_run("    "+"雨：") 
run.font.name = u'宋体'
run.font.size = Pt(12)
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(255,0,0)
run=p12.add_run(rain) 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(0,0,0)
#风
p13 = document.add_paragraph()
run=p13.add_run("    "+"风：") 
run.font.name = u'宋体'
run.font.size = Pt(12)
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(255,0,0)
run=p13.add_run(wind) 
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.color.rgb = RGBColor(0,0,0)
now = datetime.datetime.now()
date = str(now.year) + '年' + str(now.month) + '月' + str(now.day) + '日'
p=os.path.dirname(__file__)
p = os.path.join(p,os.path.pardir)
document.save('{}/'.format(p)+date+'气象.docx')




